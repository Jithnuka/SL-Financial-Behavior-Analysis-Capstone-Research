from docx import Document
from docx.shared import Inches
import os

def add_image_paragraph(doc, img_path, title, caption):
    doc.add_heading(title, level=2)
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph(caption)


def main():
    doc = Document()
    doc.add_heading('Interpretation of Charts — Borrowing Behaviour Analysis', level=1)

    figs = [
        ('output/figures/credit_sources.png', 'Primary Sources of Credit', 'Informal lenders most common; see notes.'),
        ('output/figures/borrow_gender.png', 'Borrowing by Gender', 'Similar borrowing rates for men and women.'),
        ('output/figures/borrow_income.png', 'Borrowing by Income Group', 'Middle quintiles show higher borrowing.'),
        ('output/figures/borrow_age.png', 'Borrowing by Age Group', 'Borrowing peaks among prime-age groups.'),
        ('output/figures/borrow_purpose.png', 'Main Purpose of Borrowing', 'Purpose labels are mapped from detected fin variables; verify with codebook.')
    ]

    for img, title, caption in figs:
        if os.path.exists(img):
            add_image_paragraph(doc, img, title, caption)
        else:
            doc.add_paragraph(f'Missing figure: {img}')

    doc.add_heading('Notes', level=2)
    doc.add_paragraph('Assumptions: fin22* mapped to borrow source flags; fin24*/fin30 mapped to purpose categories by heuristic. Confirm with the survey codebook.')

    out = 'Interpret_Chart.docx'
    doc.save(out)
    print('Saved', out)


if __name__ == '__main__':
    main()
