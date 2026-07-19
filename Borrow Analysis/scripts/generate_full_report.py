from docx import Document
from docx.shared import Inches
import os


def add_image_paragraph(doc, img_path, title, caption):
    doc.add_heading(title, level=2)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
    else:
        doc.add_paragraph(f'Missing figure: {img_path}')
    doc.add_paragraph(caption)


def main():
    doc = Document()
    doc.add_heading('Full Interpret Chart — Borrowing Behaviour Analysis', level=1)

    figs = [
        ('output/figures/credit_sources.png', 'Primary Sources of Credit', 'Interpretation: see notes.'),
        ('output/figures/borrow_gender.png', 'Borrowing by Gender', 'Compare rates across genders.'),
        ('output/figures/borrow_income_advanced.png', 'Borrowing by Income Group (Advanced)', 'Includes sample sizes and bootstrap 95% CI.'),
        ('output/figures/borrow_age.png', 'Borrowing by Age Group', 'Age-group patterns.'),
        ('output/figures/borrow_purpose.png', 'Main Purpose of Borrowing', 'Purpose categories via detected fin variables.')
    ]

    for img, title, caption in figs:
        add_image_paragraph(doc, img, title, caption)

    # Add interactive links section
    doc.add_heading('Interactive Figures', level=2)
    interactive_files = [
        ('Borrowing by Income Group (interactive)', 'output/figures/borrow_income_advanced.html')
    ]
    for label, path in interactive_files:
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.add_run(label + ': ')
            p.add_run(path)
        else:
            doc.add_paragraph(f'Missing interactive file: {path}')

    doc.add_heading('Notes and Assumptions', level=2)
    doc.add_paragraph('Assumptions: mapping of fin22*/fin24*/fin30 variables was done heuristically. Please verify with the survey codebook before publication.')

    out = 'Full_Interpret_Chart_1.docx'
    doc.save(out)
    print('Saved', out)


if __name__ == '__main__':
    main()
